import io
import json
import logging
import re
from pathlib import Path

import requests
from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Pt

from .builder import convert as convert_doc_template
from .styles import clean, style_run

log = logging.getLogger(__name__)

DEFAULT_TEMPLATE_URL = 'https://docs.google.com/document/d/1S02noSARXsm5hsHUfGO_Rpccq5td_1fe/edit'
DEFAULT_TEMPLATE_DOCX = Path(__file__).resolve().parent / 'assets' / 'exam_template.docx'
_FORM_DATA_RE  = re.compile(r'var FB_PUBLIC_LOAD_DATA_ = (.*?);</script>', re.S)
_FORM_ID_RE    = re.compile(r'/forms/d/e/([^/]+)|/forms/d/([^/]+)')
_ROMAN_SPLIT_RE = re.compile(
    r'\s{2,}((?:I{1,3}|IV|VI{0,3}|IX|X)|\d)\.\s+',
    re.IGNORECASE,
)
_ROMAN_MAP = {
    'I':1,'II':2,'III':3,'IV':4,'V':5,
    'VI':6,'VII':7,'VIII':8,'IX':9,'X':10,
}
_SKIP_TITLES = {
    'email',
    'full name (surname, first name, middle name)',
    'complete name of school',
    'branch',
    'block',
    'branch & block',
    'branch and block',
    'branch/block',
}


def is_form_url(url: str) -> bool:
    return 'docs.google.com/forms/' in url


def extract_form_id(url: str) -> str | None:
    m = _FORM_ID_RE.search(url.strip())
    if not m:
        return None
    return next((g for g in m.groups() if g), None)


def _download_html(url: str) -> str:
    log.info('Downloading Google Form from %s', url)
    r = requests.get(url, headers={'User-Agent': 'Mozilla/5.0'}, timeout=30)
    r.raise_for_status()
    return r.text


def _load_public_data(html: str) -> list:
    m = _FORM_DATA_RE.search(html)
    if not m:
        raise ValueError('No Google Forms data found. The form may be private or inaccessible.')
    return json.loads(m.group(1))


def _norm_title(text: str | None) -> str:
    return clean(text or '').strip().rstrip('*').strip()


def _is_skipped_title(title: str) -> bool:
    compact = re.sub(r'\s+', ' ', title).strip().lower()
    if compact in _SKIP_TITLES:
        return True
    return bool(re.fullmatch(r'(branch|block)(\s*[/&-]\s*|\s+and\s+)(branch|block)', compact))


def _extract_choices(item: list) -> list[str]:
    entries = item[4] if len(item) > 4 and isinstance(item[4], list) else []
    if not entries:
        return []
    raw_choices = entries[0][1] if len(entries[0]) > 1 and isinstance(entries[0][1], list) else []
    return [_norm_title(choice[0]) for choice in raw_choices if choice and _norm_title(choice[0])]


def _extract_items(data: list) -> tuple[str, list[dict]]:
    form = data[1]
    title = _norm_title(form[8] if len(form) > 8 else '') or 'Google Form'
    items = []

    for raw in form[1] or []:
        title_text = _norm_title(raw[1] if len(raw) > 1 else '')
        if not title_text or _is_skipped_title(title_text):
            continue

        item_type = raw[3] if len(raw) > 3 else None
        if item_type == 6:
            items.append({'type': 'section', 'text': title_text})
            continue

        choices = _extract_choices(raw)
        if choices:
            items.append({'type': 'question', 'text': title_text, 'choices': choices})

    return title, items


def _style_paragraph(p, *, font_size: int = 9, bold: bool = False) -> None:
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.0
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    for run in p.runs:
        style_run(run, {'ts_ff': 'Tahoma', 'ts_fs': font_size, 'ts_bd': bold})


def _title_lines(title: str) -> list[str]:
    m = re.search(r'recalls?\s*(\d+)\s+examinations?.*?\bNP\s*([IVX0-9]+)', title, re.I)
    if not m:
        m = re.search(r'recalls?\s+examinations?\s*(\d+).*?\bNP\s*([IVX0-9]+)', title, re.I)
    if not m:
        return [title]
    exam_no, practice = m.groups()
    practice = {'1': 'I', '2': 'II', '3': 'III', '4': 'IV', '5': 'V'}.get(practice, practice)
    return [
        f'RECALLS {exam_no} EXAMINATIONS',
        f'NURSING PRACTICE {practice}',
        'COMMUNITY HEALTH NURSING',
    ]


def _replace_template_title(doc, title: str) -> None:
    if len(doc.paragraphs) < 2:
        return
    p = doc.paragraphs[1]
    p.clear()
    lines = _title_lines(title)
    for i, line in enumerate(lines):
        if i:
            p.add_run().add_break()
        r = p.add_run(line)
        style_run(r, {'ts_ff': 'Tahoma', 'ts_fs': 16, 'ts_bd': True})
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE


def _first_question_paragraph_index(doc) -> int:
    for idx, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if idx > 2 and re.match(r'^(Situation[:\s]|Situation\s*\d+)', text, re.I):
            return idx
    return min(len(doc.paragraphs), 9)


def _remove_template_questions(doc) -> None:
    keep_until = _first_question_paragraph_index(doc)
    for p in list(doc.paragraphs[keep_until:]):
        p._element.getparent().remove(p._element)


def _strip_section_breaks(doc) -> None:
    from docx.oxml.ns import qn as _qn
    for p in doc.paragraphs:
        pPr = p._p.find(_qn('w:pPr'))
        if pPr is not None:
            for s in pPr.findall(_qn('w:sectPr')):
                pPr.remove(s)
    body = doc.element.body
    body_sectPr = body.find(_qn('w:sectPr'))
    if body_sectPr is not None:
        for cols in body_sectPr.findall(_qn('w:cols')):
            body_sectPr.remove(cols)


def _load_template_doc(template_url: str | None) -> Document:
    if template_url is None and DEFAULT_TEMPLATE_DOCX.exists():
        doc = Document(DEFAULT_TEMPLATE_DOCX)
        _strip_section_breaks(doc)
        return doc

    template_bytes, _ = convert_doc_template(template_url or DEFAULT_TEMPLATE_URL)
    doc = Document(io.BytesIO(template_bytes))
    _strip_section_breaks(doc)
    return doc


def _apply_frame_columns(doc) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    body_sectPr = doc.element.body.find(qn('w:sectPr'))
    pg_w, pg_mar_l, pg_mar_r = 12242, 720, 720
    if body_sectPr is not None:
        pgSz  = body_sectPr.find(qn('w:pgSz'))
        pgMar = body_sectPr.find(qn('w:pgMar'))
        if pgSz  is not None: pg_w     = int(pgSz .get(qn('w:w'),    pg_w))
        if pgMar is not None:
            pg_mar_l = int(pgMar.get(qn('w:left'),  pg_mar_l))
            pg_mar_r = int(pgMar.get(qn('w:right'), pg_mar_r))

    usable_w = pg_w - pg_mar_l - pg_mar_r

    for old in body_sectPr.findall(qn('w:cols')):
        body_sectPr.remove(old)
    cols_el = OxmlElement('w:cols')
    cols_el.set(qn('w:num'),   '2')
    cols_el.set(qn('w:space'), '720')
    docGrid = body_sectPr.find(qn('w:docGrid'))
    if docGrid is not None: docGrid.addprevious(cols_el)
    else:                   body_sectPr.append(cols_el)

    for p in doc.paragraphs:
        pPr = p._p.find(qn('w:pPr'))
        if pPr is None:
            pPr = OxmlElement('w:pPr')
            p._p.insert(0, pPr)
        for old_fp in pPr.findall(qn('w:framePr')):
            pPr.remove(old_fp)
        fp = OxmlElement('w:framePr')
        fp.set(qn('w:w'),       str(usable_w))
        fp.set(qn('w:wrap'),    'notBeside')
        fp.set(qn('w:vAnchor'), 'margin')
        fp.set(qn('w:hAnchor'), 'margin')
        fp.set(qn('w:x'),       '0')
        fp.set(qn('w:y'),       '0')
        pPr.insert(0, fp)


def _split_roman_subitems(text: str) -> tuple[str, list[tuple[int, str]]]:
    parts = _ROMAN_SPLIT_RE.split(text.strip())
    if len(parts) <= 1:
        return text.strip(), []
    stem     = parts[0].strip()
    subitems = []
    for i in range(1, len(parts) - 1, 2):
        label   = parts[i].upper()
        content = parts[i + 1].strip()
        if content:
            num = int(label) if label.isdigit() else _ROMAN_MAP.get(label, i // 2 + 1)
            subitems.append((num, content))
    return stem, subitems


def _add_section(doc, text: str) -> None:
    p = doc.add_paragraph()
    p.add_run(text)
    _style_paragraph(p, font_size=9, bold=True)


def _add_question(doc, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent       = Pt(18)
    p.paragraph_format.first_line_indent = Pt(-18)
    text = re.sub(r'^(\d{1,3}\.)([ \t\xa0]*)(\S)', r'\1\t\3', text, count=1)
    p.add_run(text)
    _style_paragraph(p, font_size=9)


def _add_subitem(doc, num: int, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent       = Pt(54)
    p.paragraph_format.first_line_indent = Pt(-18)
    p.add_run(f'{num}.\t{text}')
    _style_paragraph(p, font_size=9)


def _add_choice(doc, text: str, index: int = 0) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent       = Pt(36)
    p.paragraph_format.first_line_indent = Pt(-18)
    if re.match(r'^[A-D]\.', text.strip()):
        text = re.sub(r'^([A-D]\.)([ \t\xa0]*)(\S)', r'\1\t\3', text, count=1)
    else:
        text = f'{chr(ord("A") + index)}.\t{text}'
    p.add_run(text)
    _style_paragraph(p, font_size=9)


def build_docx(title: str, items: list[dict], template_url: str | None = None) -> bytes:
    doc = _load_template_doc(template_url)
    _replace_template_title(doc, title)
    _remove_template_questions(doc)
    _apply_frame_columns(doc)

    for item in items:
        if item['type'] == 'section':
            _add_section(doc, item['text'])
        elif item['type'] == 'question':
            stem, subitems = _split_roman_subitems(item['text'])
            _add_question(doc, stem)
            for num, sub_text in subitems:
                _add_subitem(doc, num, sub_text)
            for ci, choice in enumerate(item['choices']):
                _add_choice(doc, choice, ci)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def convert(url: str, template_url: str | None = None, status_fn=None) -> tuple[bytes, str]:
    def _status(msg):
        if status_fn:
            status_fn(msg)

    _status('Downloading form...')
    html = _download_html(url)

    _status('Processing questions...')
    data = _load_public_data(html)
    title, items = _extract_items(data)
    if not items:
        raise ValueError('No quiz items found in the Google Form.')
    log.info('Parsed Google Form: %s (%d output items)', title, len(items))

    _status('Building DOCX...')
    return build_docx(title, items, template_url=template_url), title
