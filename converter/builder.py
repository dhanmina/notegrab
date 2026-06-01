import io
import logging
import re

from docx import Document
from docx.shared import Pt, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .fetcher import extract_doc_id, download_html, fetch_image
from .parser import build_model, parse_chunks, extract_image_urls, DocModel
from .styles import ALIGN_MAP, clean, get_ts, style_run
from .xml_helpers import (
    add_border, add_tab_stops, extract_col_list,
    get_cols_at, inject_col_section_break, inline_to_anchor,
)
from .lists import make_label

log = logging.getLogger(__name__)

_PAGE_W_PT = 612.1
_MANUAL_NUMBERED_ITEM_RE = re.compile(r'^(\d{1,3}\.)([ \t\xa0]*)(\S)')
_SITUATION_RE = re.compile(r'^Situation[\s\d]', re.IGNORECASE)


def _init_document() -> Document:
    doc = Document()
    for sec in doc.sections:
        sec.page_width = Emu(int(_PAGE_W_PT * 12700))
        sec.page_height = Emu(int(936.1 * 12700))
        sec.top_margin = Pt(36)
        sec.bottom_margin = Pt(36)
        sec.left_margin = Pt(36)
        sec.right_margin = Pt(36)
        sec.header_distance = Pt(18)
    for ep in list(doc.paragraphs):
        ep._element.getparent().remove(ep._element)
    return doc


def _fetch_images(img_urls: dict[str, str]) -> dict[str, bytes]:
    images: dict[str, bytes] = {}
    for cid, url in img_urls.items():
        data = fetch_image(url)
        if data:
            images[cid] = data
    return images


def _collect_image_elements(chunks: list, images: dict[str, bytes]) -> list[dict]:
    elements = []
    for ch in chunks:
        for it in ch:
            if it.get('ty') != 'ae':
                continue
            epm = it.get('epm', {})
            eo = epm.get('ee_eo', {})
            cid = eo.get('i_cid')
            if not cid or cid not in images:
                continue
            pos = epm.get('ae_p', {})
            elements.append({
                'cid': cid,
                'w_pt': eo.get('i_wth', 100),
                'h_pt': eo.get('i_ht', 100),
                'x_pt': pos.get('p_hp', {}).get('hp_lo', 0),
                'y_pt': pos.get('p_vp', {}).get('vp_to', 0),
                'behind': pos.get('p_bd', False),
            })
    return elements


def _split_paragraphs(full_text: str, base: int) -> list[tuple[str, int | None, int]]:
    paragraphs: list[tuple[str, int | None, int]] = []
    seg = 0
    for i, ch in enumerate(full_text):
        if ch == '\n':
            paragraphs.append((full_text[seg:i], i + base, seg))
            seg = i + 1
    if seg < len(full_text):
        paragraphs.append((full_text[seg:], None, seg))

    # Drop leading empty paragraphs that exist only as logo spacers in the original.
    # After cleaning, these contain nothing visible (only private-use chars like ).
    while paragraphs and not clean(paragraphs[0][0]).strip():
        paragraphs.pop(0)

    # Drop trailing paragraphs that are empty or are the "| Page" page-number artifact
    # (Google Docs embeds a private-use glyph + "| Page" label in the body as a field).
    while paragraphs and not clean(paragraphs[-1][0]).strip().lstrip('| ').lower().replace('page', '').strip():
        paragraphs.pop()

    # Collapse runs of consecutive blank paragraphs down to one.
    # Source docs often have 2–3 blank lines between a situation header and the
    # first question, which creates a large visible gap in the output.
    result: list[tuple[str, int | None, int]] = []
    prev_blank = False
    for para in paragraphs:
        is_blank = not clean(para[0]).strip()
        if is_blank and prev_blank:
            continue
        result.append(para)
        prev_blank = is_blank

    # Drop blank paragraphs immediately before OR after a Situation header so
    # the GDocs output matches GForms (no gap around section headers).
    filtered: list[tuple[str, int | None, int]] = []
    for i, para in enumerate(result):
        if not clean(para[0]).strip():
            next_text = next(
                (clean(result[j][0]).strip() for j in range(i + 1, len(result)) if clean(result[j][0]).strip()),
                '',
            )
            if _SITUATION_RE.match(next_text):
                continue
            prev_text = next(
                (clean(result[j][0]).strip() for j in range(i - 1, -1, -1) if clean(result[j][0]).strip()),
                '',
            )
            if _SITUATION_RE.match(prev_text):
                continue
        filtered.append(para)
    return filtered


def _apply_paragraph_format(p, ps: dict, la) -> None:
    al = ps.get('ps_al')
    if al is not None and not ps.get('ps_al_i', True):
        p.alignment = ALIGN_MAP.get(al, WD_ALIGN_PARAGRAPH.LEFT)

    pf = p.paragraph_format

    # ps_ifl is an absolute position; first_line_indent in OOXML is relative to left_indent
    il = ps.get('ps_il')
    if il is not None and not ps.get('ps_il_i', True):
        pf.left_indent = Pt(il)
    ifl = ps.get('ps_ifl')
    if ifl is not None and not ps.get('ps_ifl_i', True):
        pf.first_line_indent = Pt(ifl - (il or 0))
    elif la is not None and il is not None and not ps.get('ps_il_i', True):
        pf.first_line_indent = Pt(-18)

    # Default to 0 to suppress Word's Normal-style spacing
    sb = ps.get('ps_sb')
    pf.space_before = Pt(sb) if sb is not None and not ps.get('ps_sb_i', True) else Pt(0)
    sa = ps.get('ps_sa')
    pf.space_after = Pt(sa) if sa is not None and not ps.get('ps_sa_i', True) else Pt(0)

    ls = ps.get('ps_ls')
    if ls is not None and not ps.get('ps_ls_i', True):
        pf.line_spacing = ls
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE

    brd = {k: ps[k] for k in ('ps_bt', 'ps_bb', 'ps_bl', 'ps_br') if k in ps}
    if brd:
        add_border(p, brd)

    ps_ts = ps.get('ps_ts')
    if ps_ts:
        add_tab_stops(p, ps_ts)


def _is_in_multi_col_body(pep: int | None, col_sectors: list) -> bool:
    if pep is None or not col_sectors:
        return False
    cols, _ = get_cols_at(pep, col_sectors)
    return cols > 1


def _apply_manual_question_format(p, para_text: str, pep: int | None, model: DocModel) -> str:
    if not _is_in_multi_col_body(pep, model.col_sectors):
        return para_text
    m = _MANUAL_NUMBERED_ITEM_RE.match(clean(para_text))
    if not m:
        return para_text

    pf = p.paragraph_format
    if pf.left_indent is None:
        pf.left_indent = Pt(18)
    if pf.first_line_indent is None:
        pf.first_line_indent = Pt(-18)

    # These rows are typed numbers in Google Docs, not list annotations. Replace
    # the separator with a tab so they align like real numbered-list rows.
    return _MANUAL_NUMBERED_ITEM_RE.sub(r'\1\t\3', para_text, count=1)


def _add_text_runs(p, para_text: str, str_start: int, base: int, text_anns: list) -> None:
    # Split on \x0b (soft return) BEFORE cleaning so positions stay aligned
    para_start = base + str_start
    sub_off = 0
    for seg_i, segment in enumerate(para_text.split('\x0b')):
        if seg_i > 0:
            r = p.add_run()
            r.add_break()
            sub_off += 1  # consume the \x0b
        if not segment:
            sub_off += len(segment)
            continue
        i = 0
        while i < len(segment):
            cp = base + str_start + sub_off + i
            ts = get_ts(cp, text_anns, para_start)
            j = i + 1
            while j < len(segment) and get_ts(base + str_start + sub_off + j, text_anns, para_start) == ts:
                j += 1
            run_text = clean(segment[i:j])
            if run_text:
                r = p.add_run(run_text)
                style_run(r, ts)
            i = j
        sub_off += len(segment)


def _build_paragraphs(doc: Document, paragraphs: list, model: DocModel) -> None:
    list_ctrs: dict = {}
    list_base_il: dict = {}  # ls_id -> minimum indent seen (parent level)
    for idx, (para_text, pep, str_start) in enumerate(paragraphs):
        ps = model.para_styles.get(pep, {}) if pep else {}
        la = model.list_anns.get(pep)
        p = doc.add_paragraph()

        _apply_paragraph_format(p, ps, la)

        # Normalize list-item indent: source docs sometimes place numbered
        # sub-items at question-level (18pt) instead of choice-level (36pt).
        # Promote them so all list items sit at a consistent 36pt indent.
        if la and p.paragraph_format.left_indent is not None:
            if abs(p.paragraph_format.left_indent.pt - 18) < 1:
                p.paragraph_format.left_indent       = Pt(36)
                p.paragraph_format.first_line_indent = Pt(-18)

        if la and la.get('ls_id') in model.list_defs:
            ls_id = la.get('ls_id', '')
            ps_il = ps.get('ps_il', 0)

            if ls_id not in list_base_il or ps_il < list_base_il[ls_id]:
                list_base_il[ls_id] = ps_il
            base_il = list_base_il[ls_id]

            # Returning to the parent indent: reset sub-level counters so they restart
            if abs(ps_il - base_il) < 1:
                for k in [k for k in list(list_ctrs) if k[0] == ls_id and k[2] != round(ps_il)]:
                    del list_ctrs[k]

            lbl, lts = make_label(ls_id, ps.get('ps_sm', 1), la, model.list_defs, list_ctrs, ps_il, base_il)
            lr = p.add_run(clean(lbl))
            style_run(lr, lts)
        elif not la:
            para_text = _apply_manual_question_format(p, para_text, pep, model)

        _add_text_runs(p, para_text, str_start, model.base, model.text_anns)

        if pep is not None and model.col_sectors:
            cur_cols, cur_space = get_cols_at(pep, model.col_sectors)
            next_pep = paragraphs[idx + 1][1] if idx + 1 < len(paragraphs) else None
            if next_pep is not None:
                next_cols, _ = get_cols_at(next_pep, model.col_sectors)
                if next_cols != cur_cols:
                    inject_col_section_break(p, cur_cols, cur_space)


def _apply_final_columns(doc: Document, col_sectors: list, last_pep: int | None = None) -> None:
    if not col_sectors:
        return
    if last_pep is not None:
        # Use the column count actually in effect at the last paragraph position.
        # This handles documents that end inside a multi-col section where the
        # final col_sector entry is single-col (the trailing sentinel sectors).
        n, space_tw = get_cols_at(last_pep, col_sectors)
        if n <= 1:
            return
    else:
        css = extract_col_list(col_sectors[-1][1])
        if len(css) <= 1:
            return
        n = len(css)
        space_tw = int((css[0].get('scol_pe') or 36) * 20)
    body = doc.element.body
    sectPr = body.find(qn('w:sectPr'))
    if sectPr is not None:
        for old in sectPr.findall(qn('w:cols')):
            sectPr.remove(old)
        for old in sectPr.findall(qn('w:type')):
            sectPr.remove(old)
        wcols = OxmlElement('w:cols')
        wcols.set(qn('w:num'), str(n))
        wcols.set(qn('w:space'), str(space_tw))
        # OOXML schema requires w:cols to precede w:docGrid; appending at the end
        # puts it after w:docGrid and Word silently ignores the column setting.
        docGrid = sectPr.find(qn('w:docGrid'))
        if docGrid is not None:
            docGrid.addprevious(wcols)
        else:
            sectPr.append(wcols)
        log.debug('Applied %d-column layout to body sectPr', n)


_SMALL_NUM_RE = re.compile(r'^[1-9]\.\t')


def _normalize_subitems(doc: Document) -> None:
    """Indent numbered sub-items one level deeper than A/B/C/D choices.

    Target indent hierarchy:
      li=18  →  question number  (N. Question text?)
      li=36  →  lettered choice  (A. / B. / C. / D.)
      li=54  →  numeric sub-item (1. / 2. / 3. …)

    Two passes:
      1. Manual runs: ≥2 consecutive li=18 small-number paragraphs → promote
         to 54pt.  Single-digit question numbers (Q1–Q9) stand alone, never
         appear as two consecutive small-numbered paragraphs.
      2. List-annotated items already at 36pt with numeric label → deepen to
         54pt so they match manually-formatted sub-items.
    """
    paras = doc.paragraphs
    n = len(paras)

    # Pass 1 — manual (no list annotation) sub-item runs
    i = 0
    while i < n:
        pf = paras[i].paragraph_format
        if pf.left_indent is None or abs(pf.left_indent.pt - 18) > 1:
            i += 1
            continue
        if not _SMALL_NUM_RE.match(paras[i].text.strip()):
            i += 1
            continue

        run = [i]
        j = i + 1
        while j < n:
            tj = paras[j].text.strip()
            if not tj:
                j += 1
                continue
            pfj = paras[j].paragraph_format
            if pfj.left_indent is not None and abs(pfj.left_indent.pt - 18) < 1 and _SMALL_NUM_RE.match(tj):
                run.append(j)
                j += 1
            else:
                break

        if len(run) >= 2:
            for k in run:
                paras[k].paragraph_format.left_indent       = Pt(54)
                paras[k].paragraph_format.first_line_indent = Pt(-18)
            i = j
        else:
            i += 1

    # Pass 2 — list-annotated sub-items already at 36pt with numeric label
    for p in paras:
        pf = p.paragraph_format
        if pf.left_indent is None or abs(pf.left_indent.pt - 36) > 1:
            continue
        if _SMALL_NUM_RE.match(p.text.strip()):
            pf.left_indent = Pt(54)


def _bold_situation_paragraphs(doc: Document) -> None:
    for p in doc.paragraphs:
        if _SITUATION_RE.match(p.text.strip()):
            for run in p.runs:
                run.bold = True
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)


def _fix_column_transition(doc: Document) -> None:
    """Replace 1-col→N-col continuous section breaks with a full-width framePr header.

    builder.py uses inline sectPrs (type=continuous, cols=1) to model column
    transitions inherited from Google Docs.  These are renderer-unreliable: some
    renderers start the multi-column section on a new page.

    Fix: find the first such transition, apply identical framePr to every preceding
    paragraph so they form one full-width frame above the columns, then remove the
    inline sectPr.  The body sectPr's multi-column setting is kept as-is.
    """
    body_sectPr = doc.element.body.find(qn('w:sectPr'))
    if body_sectPr is None:
        return
    body_cols_el = body_sectPr.find(qn('w:cols'))
    if body_cols_el is None or int(body_cols_el.get(qn('w:num'), '1')) <= 1:
        return  # single-column document — nothing to fix

    # Locate the first inline sectPr that is continuous and single-column
    transition_idx = None
    for i, p in enumerate(doc.paragraphs):
        pPr = p._p.find(qn('w:pPr'))
        if pPr is None:
            continue
        sp = pPr.find(qn('w:sectPr'))
        if sp is None:
            continue
        t = sp.find(qn('w:type'))
        if t is not None and t.get(qn('w:val')) != 'continuous':
            continue
        c = sp.find(qn('w:cols'))
        if c is None or int(c.get(qn('w:num'), '1')) == 1:
            transition_idx = i
            break

    if transition_idx is None:
        return

    # Compute usable page width from body sectPr geometry
    pg_w, pg_mar_l, pg_mar_r = 12242, 720, 720
    pgSz  = body_sectPr.find(qn('w:pgSz'))
    pgMar = body_sectPr.find(qn('w:pgMar'))
    if pgSz  is not None: pg_w     = int(pgSz .get(qn('w:w'),    pg_w))
    if pgMar is not None:
        pg_mar_l = int(pgMar.get(qn('w:left'),  pg_mar_l))
        pg_mar_r = int(pgMar.get(qn('w:right'), pg_mar_r))
    usable_w = pg_w - pg_mar_l - pg_mar_r

    # Apply identical framePr to all header paragraphs so they form one full-width frame
    for p in doc.paragraphs[:transition_idx + 1]:
        pPr = p._p.find(qn('w:pPr'))
        if pPr is None:
            pPr = OxmlElement('w:pPr')
            p._p.insert(0, pPr)
        for old in pPr.findall(qn('w:framePr')):
            pPr.remove(old)
        fp = OxmlElement('w:framePr')
        fp.set(qn('w:w'),       str(usable_w))
        fp.set(qn('w:wrap'),    'notBeside')
        fp.set(qn('w:vAnchor'), 'margin')
        fp.set(qn('w:hAnchor'), 'margin')
        fp.set(qn('w:x'),       '0')
        fp.set(qn('w:y'),       '0')
        pPr.insert(0, fp)

    # Remove the now-redundant inline sectPr from the transition paragraph
    pPr = doc.paragraphs[transition_idx]._p.find(qn('w:pPr'))
    if pPr is not None:
        for sp in pPr.findall(qn('w:sectPr')):
            pPr.remove(sp)

    log.debug('Column transition fixed: framePr on %d header paras, inline sectPr removed',
              transition_idx + 1)


def _prepend_logo(doc: Document, logo_imgs: list, images: dict[str, bytes]) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    for m in logo_imgs:
        run = p.add_run()
        run.add_picture(io.BytesIO(images[m['cid']]), width=Pt(m['w_pt']), height=Pt(m['h_pt']))


def _build_headers(doc: Document, img_elements: list, images: dict[str, bytes]) -> None:
    bg_imgs = [m for m in img_elements if m['y_pt'] >= 0]
    for sec in doc.sections:
        hdr = sec.header
        hdr.is_linked_to_previous = False
        hp = hdr.paragraphs[0]
        for m in bg_imgs:
            run = hp.add_run()
            run.add_picture(io.BytesIO(images[m['cid']]), width=Pt(m['w_pt']), height=Pt(m['h_pt']))
            x = (_PAGE_W_PT - m['w_pt']) / 2 if m['w_pt'] < _PAGE_W_PT else 0
            inline_to_anchor(run, x_pt=x, y_pt=m['y_pt'], behind=m['behind'], v_relative_from='page')



def _build_footer(doc: Document) -> None:
    sec = doc.sections[0]
    ftr = sec.footer
    ftr.is_linked_to_previous = False
    _populate_footer_paragraph(ftr.paragraphs[0])


def _populate_footer_paragraph(fp) -> None:
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(4)
    fp.paragraph_format.space_after = Pt(0)

    pPr = fp._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), '4')
    top.set(qn('w:space'), '1')
    top.set(qn('w:color'), 'AAAAAA')
    pBdr.append(top)
    pPr.append(pBdr)

    def _make_run(color='000000'):
        r = OxmlElement('w:r')
        rpr = OxmlElement('w:rPr')
        fonts = OxmlElement('w:rFonts')
        fonts.set(qn('w:ascii'), 'Arial')
        fonts.set(qn('w:hAnsi'), 'Arial')
        rpr.append(fonts)
        if color:
            c = OxmlElement('w:color')
            c.set(qn('w:val'), color)
            rpr.append(c)
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), '20')
        rpr.append(sz)
        r.append(rpr)
        return r

    # begin
    r_begin = _make_run()
    fc_begin = OxmlElement('w:fldChar')
    fc_begin.set(qn('w:fldCharType'), 'begin')
    r_begin.append(fc_begin)
    fp._p.append(r_begin)

    # instr
    r_instr = _make_run()
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' PAGE '
    r_instr.append(instr)
    fp._p.append(r_instr)

    # separate + display value (Word replaces "1" with the real page number)
    r_sep = _make_run()
    fc_sep = OxmlElement('w:fldChar')
    fc_sep.set(qn('w:fldCharType'), 'separate')
    r_sep.append(fc_sep)
    fp._p.append(r_sep)

    r_val = _make_run()
    t = OxmlElement('w:t')
    t.text = '1'
    r_val.append(t)
    fp._p.append(r_val)

    # end
    r_end = _make_run()
    fc_end = OxmlElement('w:fldChar')
    fc_end.set(qn('w:fldCharType'), 'end')
    r_end.append(fc_end)
    fp._p.append(r_end)

    fr_pipe = fp.add_run(' | ')
    fr_pipe.font.name = 'Arial'
    fr_pipe.font.size = Pt(10)
    fr_pipe.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    fr_label = fp.add_run('Page')
    fr_label.font.name = 'Arial'
    fr_label.font.size = Pt(10)
    fr_label.font.color.rgb = RGBColor(0x80, 0x80, 0x80)


def _extract_title(html: str) -> str:
    m = re.search(r'<title>([^<]+)</title>', html, re.IGNORECASE)
    if not m:
        return ""
    title = m.group(1)
    for suffix in [" - Google Docs", " - Google Документи"]:
        if title.endswith(suffix):
            return title[:-len(suffix)].strip()
    return title.strip()


def convert(url_or_id: str, status_fn=None) -> tuple[bytes, str]:
    def _status(msg):
        if status_fn:
            status_fn(msg)

    log.info('Starting conversion for: %s', url_or_id)

    doc_id = extract_doc_id(url_or_id)
    if not doc_id:
        raise ValueError('Could not extract a document ID from the provided URL.')

    _status('Downloading document...')
    html = download_html(doc_id)
    log.info('Downloaded HTML (%d bytes)', len(html))

    title = _extract_title(html)

    _status('Parsing content...')
    chunks = parse_chunks(html)
    if not chunks:
        raise ValueError('No document model found. The document may be private or inaccessible.')
    model = build_model(chunks)
    log.info('Built model: %d chars, %d para styles', len(model.full_text), len(model.para_styles))

    img_urls = extract_image_urls(html)
    if img_urls:
        _status(f'Fetching images ({len(img_urls)})...')
    images = _fetch_images(img_urls)

    img_elements = _collect_image_elements(chunks, images)
    logo_imgs = [m for m in img_elements if m['y_pt'] < 0]
    bg_imgs = [m for m in img_elements if m['y_pt'] >= 0]

    _status('Building DOCX...')
    doc = _init_document()
    paragraphs = _split_paragraphs(model.full_text, model.base)

    if logo_imgs:
        _prepend_logo(doc, logo_imgs, images)

    _build_paragraphs(doc, paragraphs, model)
    log.info('Built %d document paragraphs', len(doc.paragraphs))

    last_pep = next((pep for _, pep, _ in reversed(paragraphs) if pep is not None), None)
    _apply_final_columns(doc, model.col_sectors, last_pep)
    _normalize_subitems(doc)
    _bold_situation_paragraphs(doc)
    _fix_column_transition(doc)

    if img_elements:
        _build_headers(doc, img_elements, images)
        log.info('Built headers (%d background image(s))', len(bg_imgs))

    _build_footer(doc)
    log.info('Built footer')

    buf = io.BytesIO()
    doc.save(buf)
    size = buf.tell()
    log.info('Saved document (%d bytes)', size)
    buf.seek(0)
    return buf.read(), title
